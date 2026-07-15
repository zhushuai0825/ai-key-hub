#!/usr/bin/env python3
"""短剧剧本 Markdown → Word 导出脚本（纯 python-docx 版）

用法: python3 scripts/export_docx.py <输入.md> <输出.docx>

功能:
  1. 检测 python-docx / lxml 是否安装，缺失时自动通过 pip 安装
  2. 逐行解析剧本 Markdown，清理 Markdown 标记与内部创作骨架
  3. 用 python-docx 直接生成 .docx，精确控制中文字体（w:eastAsia）和间距
  4. 参照《女相师》类行业交付稿：A4、宋体 12pt、全文加粗、1.5倍行距

注意: 第三个位置参数（reference-doc 路径）保留接口兼容性；当前内置行业交付稿版式。
"""

import re
import subprocess
import sys
from pathlib import Path


# ─────────────────────── 依赖检测与自动安装 ───────────────────────

def _pip_install(pkgs: list[str]) -> bool:
    """尝试用 pip 安装缺失包，返回是否成功"""
    pkg_str = " ".join(pkgs)
    print(f"[安装] 正在自动安装依赖: pip install {pkg_str}")
    print("[说明] python-docx 和 lxml 是开源 Python 库，仅需安装一次")
    try:
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "--quiet", *pkgs],
            check=True,
        )
        print(f"[成功] 安装完成: {pkg_str}")
        return True
    except subprocess.CalledProcessError:
        return False


def check_deps() -> bool:
    """检测依赖，缺失时自动安装；返回最终是否可用"""
    missing = []
    try:
        import docx  # noqa
    except ImportError:
        missing.append("python-docx")
    try:
        from lxml import etree  # noqa
    except ImportError:
        missing.append("lxml")

    if not missing:
        return True

    # 尝试自动安装
    if _pip_install(missing):
        # 重新验证
        ok = True
        if "python-docx" in missing:
            try:
                import docx  # noqa
            except ImportError:
                ok = False
        if "lxml" in missing:
            try:
                from lxml import etree  # noqa
            except ImportError:
                ok = False
        if ok:
            return True

    # 自动安装失败，给出手动提示
    pkgs = " ".join(missing)
    print(f"[错误] 自动安装失败，请手动安装后重试:", file=sys.stderr)
    print(f"  pip3 install {pkgs}", file=sys.stderr)
    return False


# ─────────────────────── 文档初始化 ───────────────────────

def _set_east_asia(rpr_elem, font_cn: str):
    """在 rPr XML 元素上设置 w:eastAsia 中文字体属性"""
    from docx.oxml.ns import qn
    from lxml import etree

    rFonts = rpr_elem.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rpr_elem, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_cn)


def init_document():
    """创建文档，设置 A4 页面 + 《女相师》类行业交付稿样式"""
    from docx import Document
    from docx.shared import Pt, Cm, Inches

    doc = Document()

    # A4 页面；左右 1.25 inch、上下 1 inch，与参考稿的正文宽度一致。
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

    # Normal 基础样式：宋体 12pt 加粗，1.5倍行距，参考稿式段距。
    normal = doc.styles["Normal"]
    normal.font.size  = Pt(12)
    normal.font.bold  = True
    normal.font.name  = "宋体"
    normal.paragraph_format.line_spacing  = 1.5
    normal.paragraph_format.space_before  = Pt(12)
    normal.paragraph_format.space_after   = Pt(12)
    _set_east_asia(normal.element.get_or_add_rPr(), "宋体")

    return doc


# ─────────────────────── Run 辅助 ───────────────────────

def _make_run(para, text: str, bold: bool = True, size_pt: int = 12,
              font_en: str = "Times New Roman", font_cn: str = "宋体"):
    """添加一个 run，同时设置中英文字体"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree

    run = para.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(size_pt)
    run.font.name  = font_en

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_cn)
    return run


def _add_para(doc, text: str, *,
              bold: bool = True, size_pt: int = 12,
              font_en: str = "宋体", font_cn: str = "宋体",
              space_before: float = 12, space_after: float = 12,
              alignment=None):
    """
    添加段落，自动解析内联 **粗体** 标记。
    ** 包裹的部分强制加粗，其余部分使用 bold 参数值。
    """
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if alignment is not None:
        p.paragraph_format.alignment = alignment

    # 拆分内联 **bold** 标记（非贪婪，不跨行）
    segments = re.split(r"(\*\*[^*\n]+\*\*)", text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            _make_run(p, seg[2:-2], bold=True,
                      size_pt=size_pt, font_en=font_en, font_cn=font_cn)
        else:
            _make_run(p, seg, bold=bold,
                      size_pt=size_pt, font_en=font_en, font_cn=font_cn)
    return p


def _add_hr(doc):
    """添加场景分隔线：底边框细线"""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")       # 0.5pt 细线
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")  # 浅灰
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────── 行分类 ───────────────────────

# 优先级从高到低匹配。导出 Word 不再使用大号标题层级，分类只用于清理内容。
_KIND_RULES = [
    (re.compile(r"^# (?!#)"),    "h1"),         # 剧本标题
    (re.compile(r"^### "),       "h3"),          # 集标题
    (re.compile(r"^## "),        "h2"),          # 场景标题
    (re.compile(r"^---\s*$"),    "hr"),          # 分隔线
    (re.compile(r"^>"),          "blockquote"),  # > 前情提要 / 关键词 等注释行
    (re.compile(r"^<!--"),       "comment"),     # <!-- HTML 注释，剧本内部标记 -->
    (re.compile(r"^△"),          "action"),      # △ 动作/场景描写
    (re.compile(r"^\*\*\S"),     "dialogue"),    # **角色名** 对白
    (re.compile(r"^（"),          "cue"),         # （BGM / 字幕 / 音效）
    (re.compile(r"^【"),          "cue"),         # 【闪回】【闪出】
]

_SECTION_TITLE_MAP = {
    "一、故事梗概": ("prefix", "剧情介绍："),
    "故事梗概": ("prefix", "剧情介绍："),
    "剧情介绍": ("prefix", "剧情介绍："),
    "二、剧情脉络": ("heading", "剧情脉络："),
    "剧情脉络": ("heading", "剧情脉络："),
    "二、人物小传": ("heading", "人物介绍"),
    "人物小传": ("heading", "人物介绍"),
    "人物介绍": ("heading", "人物介绍"),
    "三、人物小传": ("heading", "人物介绍"),
    "三、正文": ("skip", ""),
    "正文": ("skip", ""),
    "四、正文": ("skip", ""),
}

_FRONT_MATTER_LABELS = (
    "分集定位",
    "本集骨架",
)

_BODY_START_RE = re.compile(
    r"^(?:△|【|（|#{1,3}\s+|\d+\s*[-－]\s*\d+|第[一二三四五六七八九十\d]+\s*场|"
    r"\*\*[^*：:]{1,20}\*\*(?:（[^）]+）)?[：:])"
)


def _strip_inline_markdown(text: str) -> str:
    """清理会直接露在 Word 里的轻量 Markdown 标记。"""
    s = text.strip()
    if s.startswith(">"):
        s = re.sub(r"^>\s*", "", s)
    s = re.sub(r"^[-*+]\s+", "", s)
    s = re.sub(r"^#+\s*", "", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = s.replace("**", "")
    return s.strip()


def _normalize_episode_title(text: str) -> str:
    """把“第1集”统一成参考稿更常见的中文数字标题。"""
    digits = "零一二三四五六七八九"

    def convert_under_100(num: int) -> str:
        if num < 10:
            return digits[num]
        if num == 10:
            return "十"
        if num < 20:
            return "十" + digits[num % 10]
        tens, ones = divmod(num, 10)
        return digits[tens] + "十" + (digits[ones] if ones else "")

    def repl(match):
        return f"第{convert_under_100(int(match.group(1)))}集"

    return re.sub(r"第\s*(\d{1,2})\s*集", repl, text)


def _iter_export_lines(md_text: str):
    """
    输出适合 Word 交付稿的行。

    - 去掉 fenced code 标记，只保留代码块内容。
    - 单集导出的开头创作骨架（分集定位 / 本集骨架）默认剥离，避免内部工作笔记进交付稿。
    - 将完整导出的章节名映射为《女相师》类命名。
    """
    in_code = False
    skipping_front_matter = False
    pending_prefix = ""

    for raw in md_text.splitlines():
        s = raw.strip()

        if s.startswith("```"):
            in_code = not in_code
            continue

        if not s:
            skipping_front_matter = False if not skipping_front_matter else skipping_front_matter
            yield ""
            continue

        cleaned = _strip_inline_markdown(s)
        if not cleaned:
            continue

        label = cleaned.rstrip("：:")
        if any(label.startswith(item) for item in _FRONT_MATTER_LABELS):
            skipping_front_matter = True
            continue
        if skipping_front_matter:
            if s == "---":
                skipping_front_matter = False
                continue
            if _BODY_START_RE.match(s):
                skipping_front_matter = False
            elif s.startswith(("-", "+")) or re.match(r"^\*\s+", s):
                continue
            else:
                skipping_front_matter = False

        if not in_code:
            section = _SECTION_TITLE_MAP.get(label)
            if section:
                mode, value = section
                if mode == "prefix":
                    pending_prefix = value
                elif mode == "heading":
                    yield value
                continue

        cleaned = _normalize_episode_title(cleaned)
        if pending_prefix:
            cleaned = pending_prefix + cleaned
            pending_prefix = ""
        yield cleaned


def classify(line: str):
    """返回 (kind, content)；content 对标题类去掉 # 前缀，blockquote 去掉 > 前缀"""
    s = line.strip()
    if not s:
        return "blank", ""
    for pattern, kind in _KIND_RULES:
        if pattern.match(s):
            if kind in ("h1", "h2", "h3"):
                content = re.sub(r"^#+\s*", "", s)
            elif kind == "blockquote":
                content = re.sub(r"^>\s*", "", s)  # 剥掉 > 前缀
            else:
                content = s
            return kind, content
    return "body", s


# ─────────────────────── 主转换 ───────────────────────

def convert(md_text: str, output_path: str) -> bool:
    doc = init_document()

    for line in _iter_export_lines(md_text):
        kind, content = classify(line)

        if kind == "blank":
            continue

        elif kind == "hr":
            # 参考稿不使用横线分隔。空白由上下段距承担。
            continue

        elif kind == "comment":
            # <!-- ... --> HTML 注释：直接跳过，不写入 Word
            continue

        else:
            _add_para(doc, content, space_before=12, space_after=12)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[完成] Word 文件已生成: {output_path}")
    return True


# ─────────────────────── 入口 ───────────────────────

def main():
    if len(sys.argv) == 2 and sys.argv[1] in {"-h", "--help"}:
        print("用法: python3 export_docx.py <输入.md> <输出.docx>")
        sys.exit(0)

    if len(sys.argv) < 3:
        print("用法: python3 export_docx.py <输入.md> <输出.docx>")
        sys.exit(1)

    if not check_deps():
        print("[提示] Mac:   pip3 install python-docx lxml")
        print("[提示] Win:   pip install python-docx lxml")
        print("[提示] Linux: pip3 install python-docx lxml")
        sys.exit(1)

    input_path  = Path(sys.argv[1])
    output_path = sys.argv[2]
    # sys.argv[3]（旧版 reference-doc 路径）已不再使用，静默忽略

    if not input_path.exists():
        print(f"[错误] 输入文件不存在: {input_path}", file=sys.stderr)
        sys.exit(1)

    md_text = input_path.read_text(encoding="utf-8")
    if not convert(md_text, output_path):
        sys.exit(1)


if __name__ == "__main__":
    main()
